import pandas as pd
from docx import Document


def subtables(df):
    columns_of_interest = ['LBSCAT', 'QSCAT', 'VSTEST']
    distinct_values = []

    endheader_index = df[df.iloc[:, 0].str.contains('Significant Digits', na=False)].index[0]
    postheader_index = endheader_index + 1
    df = df.iloc[postheader_index:]

    for column in columns_of_interest:
        # Check if column exists in the DataFrame
        if column in df.columns:
            distinct_values.extend([column] + df[column].dropna().unique().tolist())

    return distinct_values #should return list of variable and possible subcats


def transpose_select(df):
    # Transpose the sheet
    df_transposed = df.T

    # Add a new column with the row names - Preserve variables
    df_transposed.insert(0, 'Row_Name', df_transposed.index)

    # Rename the columns for clarity
    df_transposed.columns = df_transposed.iloc[0]
    df_transposed = df_transposed.iloc[1:]

    # Rename the columns to match the required names
    df_transposed = df_transposed.rename(columns={
        '[Table Name]': 'Variable',
        'Variable Label': 'Label',
        'CDISC Notes': 'Comment'
    })

    # Retain only the necessary columns
    df_transposed = df_transposed[['Variable', 'Label', 'Type', 'Comment']]

    df_transposed = df_transposed[~df_transposed.apply(lambda row: row.str.startswith('Unnamed')).any(axis=1)]

    # Step 4: Reset index
    df = df.reset_index(drop=True)

    return df_transposed


def consolidate_sheets(file_paths):
    # Dictionary to hold the consolidated data for each sheet
    consolidated_data = {}
    consolidated_subtables = {}

    # Step 1: Iterate over all file paths
    for file_path in file_paths:
        # Read all sheets from the current Excel file
        sheets = pd.read_excel(file_path, sheet_name=None)  # sheet_name=None loads all sheets as a dictionary

        # Extract the book name (file name) without the extension
        book_name = file_path.split('/')[-1].split('.')[0]

        # Step 2: Process each sheet in the current book
        excluded_sheets = ['ReadMe', 'Variables', 'Datasets', 'DR']
        has_subcategories = ['LB', 'QS', 'VS']
        for sheet_name, df in sheets.items():
            if sheet_name not in excluded_sheets:

                if sheet_name in has_subcategories:
                    sub = subtables(df) #this is a vector of a variable with distinct possible values
                    specialvariable = sub[0]
                    specialvariableoptions = sub[1:]
                    subdf = pd.DataFrame({'Subvariable': specialvariableoptions})
                    subdf['Study'] = book_name

                    # Step 3 1/2: If the subtable list is not in the consolidated subtables, add it
                    if specialvariable not in consolidated_subtables:
                        consolidated_subtables[specialvariable] = subdf
                    else:
                        # Otherwise, concatenate the current sheet's data to the existing consolidated data
                        consolidated_subtables[specialvariable] = pd.concat([consolidated_subtables[specialvariable], subdf], ignore_index=True)

                df_transposed = transpose_select(df)
                df = df_transposed

                # Add a column indicating the source book
                df['Study'] = book_name

                # Step 3: If the sheet is not in the consolidated_data dictionary, add it
                if sheet_name not in consolidated_data:
                    consolidated_data[sheet_name] = df
                else:
                    # Otherwise, concatenate the current sheet's data to the existing consolidated data
                    consolidated_data[sheet_name] = pd.concat([consolidated_data[sheet_name], df], ignore_index=True)

    # Step 4: For each sheet, perform the necessary row consolidation
    for sheet_name, df in consolidated_data.items():

        # Remove duplicates and aggregate the 'Variable' column
        consolidated_data[sheet_name] = df.groupby(['Variable'], as_index=False).agg(
            lambda x: ', '.join(sorted(x.unique())) if isinstance(x.iloc[0], str) else x.iloc[0]
        )

    # Step 4 1/2: For each special variable, perform the necessary row consolidation
    for specialvariable, subdf in consolidated_subtables.items():
        # Remove duplicates and aggregate the 'Subvariable' column
        consolidated_subtables[specialvariable] = subdf.groupby(['Subvariable'], as_index=False).agg(
            lambda x: ', '.join(sorted(x.unique())) if isinstance(x.iloc[0], str) else x.iloc[0]
        )

    # Step 5: Return the consolidated data
    return consolidated_data, consolidated_subtables


def tables(consolidated_data, doc):
    for sheet_name, df in consolidated_data.items():
        # Add a heading for the table (sheet name)
        doc.add_heading(sheet_name, level=1)

        # Add a table for the current sheet data
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'

        # Add the header row
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = column_name

        # Add the data rows
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        # Add a page break after each table (optional)
        doc.add_page_break()


def write_to_word(normal_data, sub_data, output_path):
    # Create a new Document
    doc = Document()
    doc.add_heading('Data Dictionary for CF Studies', 0)

    # Iterate through the consolidated data and add tables to the Word document
    tables(normal_data, doc)
    tables(sub_data, doc)

    # Save the document
    doc.save(output_path)


def main():

    file_paths = ['BPCFRD.xlsx', 'STRONG.xlsx']
    consolidated_data, consolidated_sub = consolidate_sheets(file_paths)

    # Write the consolidated tables to a Word document
    output_path = 'consolidated_data.docx'
    write_to_word(consolidated_data, consolidated_sub, output_path)


if __name__ == "__main__":
    main()
